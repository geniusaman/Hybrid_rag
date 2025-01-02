# requirements.txt content:
# streamlit
# langchain
# sentence-transformers
# numpy
# pypdf
# python-docx
# pandas
# openpyxl
# unstructured
# torch
# transformers
# faiss-cpu

import streamlit as st
import os
import networkx as nx
import matplotlib.pyplot as plt
import numpy as np
from dataclasses import dataclass
from sentence_transformers import util
import torch
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI, AzureChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from pypdf import PdfReader
from docx import Document
import pandas as pd
import tempfile
from typing import List, Dict, Tuple
from langchain.schema import Document
import pickle
from langchain.schema import BaseRetriever
from langchain.schema.messages import AIMessage, HumanMessage, SystemMessage
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
load_dotenv()

AZURE_OPENAI_API_KEY = os.environ["AZURE_OPENAI_API_KEY"]
AZURE_OPENAI_ENDPOINT = os.environ["AZURE_OPENAI_ENDPOINT"]

@dataclass
class GraphNode:
    """Represents a node in the knowledge graph"""
    id: str
    content: str
    embedding: np.ndarray
    metadata: Dict

class HybridKnowledgeBase:
    """Combines vector store and knowledge graph approaches"""
    def __init__(self, embedding_model):
        self.embedding_model = embedding_model
        self.vector_store = None
        self.knowledge_graph = nx.Graph()
        self.node_embeddings = {}
        self.similarity_threshold = 0.7

    def save(self, path="hybrid_kb.pkl"):
        """Save the hybrid knowledge base"""
        state = {
            'vector_store': self.vector_store,
            'knowledge_graph': self.knowledge_graph,
            'node_embeddings': self.node_embeddings
        }
        with open(path, "wb") as f:
            pickle.dump(state, f)

    @classmethod
    def load(cls, path="hybrid_kb.pkl", embedding_model=None):
        """Load the hybrid knowledge base"""
        with open(path, "rb") as f:
            state = pickle.load(f)
        
        kb = cls(embedding_model)
        kb.vector_store = state['vector_store']
        kb.knowledge_graph = state['knowledge_graph']
        kb.node_embeddings = state['node_embeddings']
        return kb

    def as_retriever(self, search_kwargs=None):
        """Create a retriever interface for the hybrid knowledge base"""
        if search_kwargs is None:
            search_kwargs = {"k": 3}

        def hybrid_retriever(query_str):
            return self.query(query_str, k=search_kwargs.get("k", 3))

        # Make the function callable like a regular retriever
        return type('HybridRetriever', (), {
            '__call__': hybrid_retriever,
            'get_relevant_documents': hybrid_retriever
        })()

    def add_documents(self, texts: List[Document]) -> None:
        """Process documents and create both vector store and knowledge graph"""
        # Create vector store
        self.vector_store = FAISS.from_documents(texts, self.embedding_model)
        
        # Create nodes for each text chunk
        nodes = []
        for i, doc in enumerate(texts):
            embedding = self.embedding_model.embed_query(doc.page_content)
            node = GraphNode(
                id=f"node_{i}",
                content=doc.page_content,
                embedding=embedding,
                metadata=doc.metadata
            )
            nodes.append(node)
            self.node_embeddings[node.id] = embedding
            self.knowledge_graph.add_node(node.id, content=doc.page_content, metadata=doc.metadata)

        # Create edges between semantically similar nodes
        self._create_semantic_edges(nodes)
        self._create_structural_edges(texts)

    def _create_semantic_edges(self, nodes: List[GraphNode]) -> None:
        """Create edges between semantically similar nodes"""
        for i, node1 in enumerate(nodes):
            for node2 in nodes[i+1:]:
                similarity = util.cos_sim(
                    torch.tensor(node1.embedding),
                    torch.tensor(node2.embedding)
                ).item()
                
                if similarity > self.similarity_threshold:
                    self.knowledge_graph.add_edge(
                        node1.id,
                        node2.id,
                        weight=similarity,
                        type='semantic'
                    )

    def _create_structural_edges(self, texts: List[Document]) -> None:
        """Create edges based on document structure"""
        current_doc = None
        prev_node = None
        
        for i, doc in enumerate(texts):
            node_id = f"node_{i}"
            
            # Connect sequential chunks from same document
            if doc.metadata['source'] == current_doc and prev_node:
                self.knowledge_graph.add_edge(
                    prev_node,
                    node_id,
                    weight=1.0,
                    type='sequential'
                )
            
            current_doc = doc.metadata['source']
            prev_node = node_id

    def query(self, query: str, k: int = 3) -> Tuple[List[Document], List[str]]:
        """Hybrid search combining vector and graph-based retrieval"""
        # Vector search
        vector_results = self.vector_store.similarity_search(query, k=k)
        
        # Graph-based search
        query_embedding = self.embedding_model.embed_query(query)
        initial_nodes = self._get_initial_nodes(query_embedding, k)
        graph_results = self._graph_search(initial_nodes, k)
        
        # Combine and deduplicate results
        combined_results = self._merge_results(vector_results, graph_results)
        return combined_results

    def _get_initial_nodes(self, query_embedding: np.ndarray, k: int) -> List[str]:
        """Get initial nodes for graph traversal based on embedding similarity"""
        similarities = {}
        for node_id, embedding in self.node_embeddings.items():
            similarity = util.cos_sim(
                torch.tensor(query_embedding),
                torch.tensor(embedding)
            ).item()
            similarities[node_id] = similarity
        
        return sorted(similarities.items(), key=lambda x: x[1], reverse=True)[:k]

    def _graph_search(self, initial_nodes: List[Tuple[str, float]], k: int) -> List[Document]:
        """Perform graph-based search starting from initial nodes"""
        results = []
        visited = set()
        
        for node_id, _ in initial_nodes:
            if len(results) >= k:
                break
                
            if node_id not in visited:
                visited.add(node_id)
                
                # Get subgraph of connected nodes
                connected = nx.single_source_shortest_path_length(
                    self.knowledge_graph,
                    node_id,
                    cutoff=2  # Limit path length
                )
                
                for connected_node in connected:
                    if connected_node not in visited:
                        visited.add(connected_node)
                        node_data = self.knowledge_graph.nodes[connected_node]
                        results.append(Document(
                            page_content=node_data['content'],
                            metadata=node_data['metadata']
                        ))
                        
                        if len(results) >= k:
                            break
        
        return results

    def _merge_results(self, vector_results: List[Document], graph_results: List[Document]) -> List[Document]:
        """Merge and deduplicate results from both approaches"""
        seen_content = set()
        merged = []
        
        for doc in vector_results + graph_results:
            if doc.page_content not in seen_content:
                seen_content.add(doc.page_content)
                merged.append(doc)
        
        return merged
    
class DocumentProcessor:
    """Handles different document types and converts them to text"""
    
    @staticmethod
    def process_pdf(file) -> str:
        """Extract text from PDF files"""
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text

    @staticmethod
    def process_docx(file) -> str:
        """Extract text from DOCX files"""
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    @staticmethod
    def process_excel(file) -> str:
        """Extract text from Excel files"""
        df = pd.read_excel(file, sheet_name=None)
        text = ""
        for sheet_name, sheet_df in df.items():
            text += f"\nSheet: {sheet_name}\n"
            text += sheet_df.to_string(index=False) + "\n"
        return text

    @classmethod
    def process_file(cls, file) -> str:
        """Process file based on its type"""
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_file_path = tmp_file.name

        try:
            file_extension = file.name.split('.')[-1].lower()
            if file_extension == 'pdf':
                text = cls.process_pdf(tmp_file_path)
            elif file_extension == 'docx':
                text = cls.process_docx(tmp_file_path)
            elif file_extension in ['xlsx', 'xls']:
                text = cls.process_excel(tmp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            return text
        finally:
            os.unlink(tmp_file_path)

def process_documents(raw_documents: List[str], filenames: List[str]) -> List[Document]:
    """Process and split documents for ingestion"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    
    documents = []
    for text, filename in zip(raw_documents, filenames):
        doc = Document(
            page_content=text,
            metadata={"source": filename}
        )
        documents.append(doc)
    
    texts = text_splitter.split_documents(documents)
    return texts

def create_vectorstore(texts):
    """Create FAISS vector store"""
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )
    
    vectorstore = FAISS.from_documents(texts, embeddings)
    return vectorstore

def save_vectorstore(vectorstore, path="vectorstore.pkl"):
    """Save the vector store to disk"""
    with open(path, "wb") as f:
        pickle.dump(vectorstore, f)

def load_vectorstore(path="vectorstore.pkl"):
    """Load the vector store from disk"""
    with open(path, "rb") as f:
        vectorstore = pickle.load(f)
    return vectorstore

class HybridRetriever(BaseRetriever):
    """Custom retriever that combines vector and graph-based search"""
    
    def __init__(self, hybrid_kb):
        """Initialize the hybrid retriever with a knowledge base"""
        # Call parent class constructor
        super().__init__()
        # Store hybrid_kb as instance variable
        self._hybrid_kb = hybrid_kb
    
    def _get_relevant_documents(self, query: str) -> List[Document]:
        """Implementation of abstract method from BaseRetriever"""
        return self._hybrid_kb.query(query)
    
    async def _aget_relevant_documents(self, query: str) -> List[Document]:
        """Async implementation of abstract method"""
        raise NotImplementedError("Async retrieval not implemented")

class CustomHybridChain:
    """Custom conversation chain implementation for hybrid RAG"""
    
    def __init__(self, llm, retriever: HybridRetriever, memory=None):
        self.llm = llm
        self.retriever = retriever
        self.memory = memory or ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
    
    def _format_chat_history(self, chat_history) -> str:
        """Format chat history for context"""
        formatted_history = ""
        
        # Handle different chat history formats
        if isinstance(chat_history, list):
            for entry in chat_history:
                # If entry is a Message object
                if isinstance(entry, (HumanMessage, AIMessage)):
                    formatted_history += f"{'Human' if isinstance(entry, HumanMessage) else 'Assistant'}: {entry.content}\n"
                # If entry is a tuple of messages
                elif isinstance(entry, tuple):
                    # Extract just the message content, ignore sources
                    if len(entry) >= 2:  # Handle tuples of any length >= 2
                        question, answer = entry[0], entry[1]
                        formatted_history += f"Human: {question}\nAssistant: {answer}\n"
        
        return formatted_history
    
    def _create_prompt(self, question: str, retrieved_docs: List[Document], chat_history: Optional[List] = None) -> str:
        """Create a prompt combining context and question"""
        # Format document context
        context = "\n\n".join([doc.page_content for doc in retrieved_docs])
        
        # Format chat history if available
        history_text = ""
        if chat_history:
            history_text = "\nChat History:\n" + self._format_chat_history(chat_history)
        
        # Construct the full prompt
#         prompt = f"""Using the following context and chat history, answer the question. If you cannot find the answer in the context, First Attempt to figure out the underlying logic within the content and generate a logical solution, If you still cannot figure it out, Then say so.

# Context:
# {context}
# {history_text}

# Question: {question}

# Answer:"""
        prompt = f"""You are having a natural conversation with a human. Incorporate relevant information from this context to help inform your responses, while maintaining a warm, conversational tone:

   Context:
  {context}
  conversation history:
  {history_text}

 Question: {question}

Remember to:
- Respond naturally as you would in a flowing conversation
- Avoid phrases like "based on the context" or "according to"
- Connect ideas smoothly without explicitly referencing the context
- Use a friendly, engaging tone
"""
        return prompt

    def invoke(self, query_dict: Dict[str, Any]) -> Dict[str, Any]:
        """Process a query and return response"""
        question = query_dict["question"]
        
        # Get chat history from memory
        chat_history = None
        if self.memory:
            memory_vars = self.memory.load_memory_variables({})
            chat_history = memory_vars.get("chat_history", [])
        
        # Retrieve relevant documents
        retrieved_docs = self.retriever.get_relevant_documents(question)
        
        # Create the prompt
        prompt = self._create_prompt(question, retrieved_docs, chat_history)
        
        # Get response from LLM
        messages = [
            SystemMessage(content="You are a helpful assistant that answers questions based on the provided context."),
            HumanMessage(content=prompt)
        ]
        response = self.llm.invoke(messages)
        
        # Update memory with just the question and answer
        if self.memory:
            self.memory.chat_memory.add_user_message(question)
            self.memory.chat_memory.add_ai_message(response.content)
        
        # Return formatted response
        return {
            "answer": response.content,
            "source_documents": retrieved_docs,
            "chat_history": chat_history
        }
    
def create_conversation_chain(hybrid_kb):
    """Create custom conversation chain using hybrid knowledge base"""
    llm = AzureChatOpenAI(
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        openai_api_key=AZURE_OPENAI_API_KEY,
        deployment_name="gpt-4o",
        openai_api_version="2024-08-01-preview"
    )
    
    # Create custom retriever with fixed implementation
    retriever = HybridRetriever(hybrid_kb)
    
    # Create custom chain
    chain = CustomHybridChain(
        llm=llm,
        retriever=retriever,
        memory=ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
    )
    
    return chain

def create_message_container():
    """Create a custom container for chat messages"""
    return st.container()

def display_message(is_user: bool, message: str, sources=None):
    """Display a single message in the chat"""
    if is_user:
        st.markdown(f"""
            <div style="display: flex; justify-content: flex-end; margin-bottom: 1rem;">
                <div style="background-color: #1976d2; color: white; padding: 0.75rem 1rem; 
                        border-radius: 15px; max-width: 80%; margin-left: 20%;">
                    {message}
                </div>
            </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
            <div style="display: flex; margin-bottom: 1rem;">
                <div style="background-color: #f0f2f6; padding: 0.75rem 1rem; 
                        border-radius: 15px; max-width: 80%;">
                    {message}
                </div>
            </div>
        """, unsafe_allow_html=True)
        if sources:
            with st.expander("View Sources", expanded=False):
                for i, doc in enumerate(sources):
                    st.markdown(f"**Source {i+1}** from: {doc.metadata.get('source', 'Unknown')}")
                    st.markdown(doc.page_content)
                    st.divider()
                    break

def handle_chat_history_display(chat_history):
    """Safely handle chat history display"""
    for entry in chat_history:
        try:
            if isinstance(entry, tuple):
                if len(entry) >= 3:
                    question, answer, sources = entry
                elif len(entry) == 2:
                    question, answer = entry
                    sources = None
                else:
                    continue
                display_message(True, question)
                display_message(False, answer, sources)
            elif isinstance(entry, (HumanMessage, AIMessage)):
                display_message(
                    isinstance(entry, HumanMessage),
                    entry.content,
                    None
                )
        except Exception as e:
            st.error(f"Error displaying message: {str(e)}")
            continue

def display_welcome_screen():
    """Display the welcome screen with logo and message"""
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
            <div style='text-align: center; padding: 2rem; margin: 2rem;'>
                <div style='background-color: #E3BBED; border-radius: 50%; width: 150px; height: 150px; 
                          margin: 0 auto; display: flex; align-items: center; justify-content: center;
                          border: 3px solid #673DB0;
                          box-shadow: 0 4px 8px rgba(103, 61, 176, 0.2);'>
                    <span style='font-size: 2rem; color: #673DB0; 
                               font-weight: 800; 
                               font-family: "Arial Black", "Helvetica", sans-serif;
                               text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
                               letter-spacing: 1px;'>
                        EaseAI
                    </span>
                </div>
                <div style='margin-top: 1rem; font-size: 1.5rem; color: #666;'>
                    Welcome to EaseworkAI
                </div>
                <div style='margin-top: 0.5rem; color: #888; font-size: 1rem;'>
                    Upload your documents and start asking questions!
                </div>
        </div>
        """, unsafe_allow_html=True)

def main():
    st.title("Hybrid RAG System with Knowledge Graph")
    
    # Initialize session state
    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'hybrid_kb' not in st.session_state:
        st.session_state.hybrid_kb = None
    
    # Move document management to sidebar
    with st.sidebar:
        st.header("Document Management")
        
        # Document upload section
        uploaded_files = st.file_uploader(
            "Upload your documents",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx', 'xlsx', 'xls']
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Load Previous Documents", use_container_width=True):
                try:
                    embeddings = HuggingFaceEmbeddings(
                        model_name="sentence-transformers/all-mpnet-base-v2",
                        model_kwargs={'device': 'cpu'},
                        encode_kwargs={'normalize_embeddings': True}
                    )
                    hybrid_kb = HybridKnowledgeBase.load(embedding_model=embeddings)
                    st.session_state.hybrid_kb = hybrid_kb
                    st.session_state.conversation = create_conversation_chain(hybrid_kb)
                    st.success("Successfully loaded previous documents!")
                except FileNotFoundError:
                    st.error("No previous documents found. Please Upload & process documents first.")
                except Exception as e:
                    st.error(f"Error loading index: {str(e)}")
        
        with col2:
            if uploaded_files and st.button("Process Documents", use_container_width=True):
                with st.spinner("Processing documents..."):
                    try:
                        # Process each document
                        processed_texts = []
                        filenames = []
                        for file in uploaded_files:
                            try:
                                text = DocumentProcessor.process_file(file)
                                processed_texts.append(text)
                                filenames.append(file.name)
                            except Exception as e:
                                st.error(f"Error processing {file.name}: {str(e)}")
                                continue
                        
                        if processed_texts:
                            texts = process_documents(processed_texts, filenames)
                            
                            # Initialize embedding model
                            embeddings = HuggingFaceEmbeddings(
                                model_name="sentence-transformers/all-mpnet-base-v2",
                                model_kwargs={'device': 'cpu'},
                                encode_kwargs={'normalize_embeddings': True}
                            )
                            
                            # Create hybrid knowledge base
                            hybrid_kb = HybridKnowledgeBase(embeddings)
                            hybrid_kb.add_documents(texts)
                            
                            # Save state
                            st.session_state.hybrid_kb = hybrid_kb
                            hybrid_kb.save()
                            
                            # Create conversation chain
                            st.session_state.conversation = create_conversation_chain(hybrid_kb)
                            
                            st.success("All documents processed and indexed successfully!")
                        else:
                            st.error("No documents were successfully processed.")
                    
                    except Exception as e:
                        st.error(f"An error occurred: {str(e)}")

        # Knowledge Graph Visualization
        if st.session_state.hybrid_kb and st.button("Visualize Knowledge Graph", use_container_width=True):
            with st.spinner("Generating graph visualization..."):
                fig, ax = plt.subplots(figsize=(10, 10))
                pos = nx.spring_layout(st.session_state.hybrid_kb.knowledge_graph)
                nx.draw(st.session_state.hybrid_kb.knowledge_graph, pos, 
                       with_labels=True, node_color='lightblue', 
                       node_size=1000, font_size=8)
                st.pyplot(fig)
        
        # Display system status
        st.sidebar.divider()
        status = "🟢 Ready" if st.session_state.conversation else "🔴 Not Ready"
        st.sidebar.markdown(f"**System Status:** {status}")

     # Main chat interface
    chat_container = st.container()
    
    # Display welcome screen if no chat history
    if not st.session_state.chat_history:
        display_welcome_screen()
    
    # Display chat history - Modified this part
    for message in st.session_state.chat_history:
        if isinstance(message, tuple) and len(message) == 3:
            question, answer, sources = message
            display_message(True, question)
            display_message(False, answer, sources)
    
    # Chat input
    st.divider()
    if prompt := st.chat_input("Ask a question about your documents"):
        if st.session_state.conversation is None:
            st.error("Please upload and process documents or load existing index first!")
        else:
            try:
                with st.spinner("Thinking..."):
                    # Get response from conversation chain
                    response = st.session_state.conversation.invoke({
                        "question": prompt
                    })
                    
                    # Add to chat history with sources
                    st.session_state.chat_history.append((
                        prompt,
                        response['answer'],
                        response.get('source_documents', [])
                    ))
                    
                    # Rerun to update the chat display
                    st.rerun()
                
            except Exception as e:
                st.error(f"Error generating response: {str(e)}")

if __name__ == "__main__":
    main()
